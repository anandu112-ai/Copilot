"""
Universal Document Ingestion Engine — CA Copilot
Handles all supported input formats and normalizes them into (text, tables, images, metadata).

Supported formats:
  PDF (digital + scanned), JPG, JPEG, PNG, TIFF, HEIC, WEBP,
  XLSX, XLS, CSV, DOCX, DOC, TXT, ZIP
"""
import io
import os
import zipfile
import tempfile
import mimetypes
from pathlib import Path
from typing import List, Tuple, Optional, Dict, Any
from loguru import logger

from processors.image_preprocessor import (
    preprocess_image_bytes,
    preprocess_pdf_pages,
    separate_invoices_in_scan,
)


# ── Mime / extension helpers ───────────────────────────────────────────────

IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".tiff", ".tif", ".bmp", ".webp", ".heic", ".heif"}
PDF_EXTS = {".pdf"}
EXCEL_EXTS = {".xlsx", ".xls"}
CSV_EXTS = {".csv"}
WORD_EXTS = {".docx", ".doc"}
TEXT_EXTS = {".txt", ".text", ".log"}
ZIP_EXTS = {".zip", ".tar", ".gz", ".rar", ".7z"}

ALL_SUPPORTED = IMAGE_EXTS | PDF_EXTS | EXCEL_EXTS | CSV_EXTS | WORD_EXTS | TEXT_EXTS | ZIP_EXTS


class IngestionResult:
    """Normalized output from any document format."""

    def __init__(self):
        self.full_text: str = ""
        self.raw_tables: List[List[List[str]]] = []
        self.page_count: int = 1
        self.is_ocr_used: bool = False
        self.image_preprocessed: bool = False
        self.source_format: str = ""
        self.file_size_bytes: int = 0
        self.embedded_images: List[bytes] = []  # For future multi-modal models


class DocumentIngestionEngine:
    def __init__(self, file_path: str, ocr_enabled: bool = True, preprocess: bool = True):
        self.file_path = file_path
        self.ocr_enabled = ocr_enabled
        self.preprocess = preprocess
        self.path = Path(file_path)
        self.ext = self.path.suffix.lower()

    async def ingest(self) -> IngestionResult:
        """Main entry — dispatch to the correct handler."""
        result = IngestionResult()
        result.source_format = self.ext.lstrip(".")
        result.file_size_bytes = self.path.stat().st_size if self.path.exists() else 0

        logger.info(f"Ingesting [{self.ext.upper()}]: {self.path.name}")

        if self.ext in PDF_EXTS:
            await self._ingest_pdf(result)
        elif self.ext in IMAGE_EXTS:
            await self._ingest_image(result)
        elif self.ext in EXCEL_EXTS:
            await self._ingest_excel(result)
        elif self.ext in CSV_EXTS:
            await self._ingest_csv(result)
        elif self.ext in WORD_EXTS:
            await self._ingest_word(result)
        elif self.ext in TEXT_EXTS:
            await self._ingest_text(result)
        elif self.ext in ZIP_EXTS:
            await self._ingest_zip(result)
        else:
            # Fallback: try text
            logger.warning(f"Unknown extension '{self.ext}' — trying text fallback")
            await self._ingest_text(result)

        logger.info(
            f"Ingestion complete: {result.page_count} pages, "
            f"OCR={result.is_ocr_used}, tables={len(result.raw_tables)}, "
            f"text_chars={len(result.full_text)}"
        )
        return result

    # ── PDF ─────────────────────────────────────────────────────────────────

    async def _ingest_pdf(self, result: IngestionResult):
        """Try native text extraction first; fall back to OCR if scanned."""
        import fitz
        import pdfplumber

        doc = fitz.open(self.file_path)
        result.page_count = len(doc)
        doc.close()

        # Check if digital PDF by measuring extractable characters
        text_preview = ""
        try:
            doc = fitz.open(self.file_path)
            for page in doc:
                text_preview += page.get_text("text")
                if len(text_preview) > 200:
                    break
            doc.close()
        except Exception:
            pass

        char_density = len(text_preview.strip()) / max(result.page_count, 1)

        if char_density > 50:
            # Digital PDF — use pdfplumber for tables + fitz for text
            logger.info("Detected digital PDF — using native text + table extraction")
            text_parts = []
            all_tables = []

            try:
                with pdfplumber.open(self.file_path) as pdf:
                    for page in pdf.pages:
                        t = page.extract_text(x_tolerance=3, y_tolerance=3) or ""
                        text_parts.append(t)
                        tables = page.extract_tables({
                            "vertical_strategy": "lines",
                            "horizontal_strategy": "lines",
                            "snap_tolerance": 5,
                            "join_tolerance": 3,
                            "edge_min_length": 3,
                            "min_words_vertical": 1,
                            "min_words_horizontal": 1,
                        })
                        for table in tables:
                            if table:
                                cleaned = [
                                    [str(c).strip() if c is not None else "" for c in row]
                                    for row in table
                                ]
                                all_tables.append(cleaned)
            except Exception as e:
                logger.warning(f"pdfplumber failed: {e} — using fitz fallback")
                doc = fitz.open(self.file_path)
                for page in doc:
                    text_parts.append(page.get_text("text"))
                doc.close()

            result.full_text = "\n".join(text_parts)
            result.raw_tables = all_tables

        else:
            # Scanned PDF — preprocess pages and run OCR
            logger.info("Detected scanned PDF — running image preprocessing + OCR")
            result.is_ocr_used = True

            if self.preprocess:
                pages_bytes = preprocess_pdf_pages(self.file_path, scale=2.5)
                result.image_preprocessed = True
            else:
                import fitz
                doc = fitz.open(self.file_path)
                mat = fitz.Matrix(2.0, 2.0)
                pages_bytes = []
                for page in doc:
                    pix = page.get_pixmap(matrix=mat, colorspace=fitz.csRGB)
                    pages_bytes.append(pix.tobytes("png"))
                doc.close()

            text_parts = []
            for i, page_bytes in enumerate(pages_bytes):
                ocr_text = _run_ocr(page_bytes)
                text_parts.append(ocr_text)
                logger.debug(f"OCR page {i+1}: {len(ocr_text)} chars")

            result.full_text = "\n\n".join(text_parts)

    # ── Image ────────────────────────────────────────────────────────────────

    async def _ingest_image(self, result: IngestionResult):
        """Preprocess image then run OCR."""
        result.is_ocr_used = True
        result.page_count = 1

        with open(self.file_path, "rb") as f:
            raw_bytes = f.read()

        # Handle HEIC conversion first
        if self.ext in {".heic", ".heif"}:
            raw_bytes = _convert_heic_to_png(raw_bytes)

        # Check for multiple invoices in single scan
        regions = separate_invoices_in_scan(raw_bytes)

        text_parts = []
        for region_bytes in regions:
            if self.preprocess:
                processed = preprocess_image_bytes(region_bytes, self.path.name)
                result.image_preprocessed = True
            else:
                processed = region_bytes
            text_parts.append(_run_ocr(processed))

        result.full_text = "\n\n--- Page Break ---\n\n".join(text_parts)
        result.embedded_images = [raw_bytes]

    # ── Excel ─────────────────────────────────────────────────────────────── 

    async def _ingest_excel(self, result: IngestionResult):
        """Read all sheets, extract text and tables."""
        import openpyxl

        wb = openpyxl.load_workbook(self.file_path, data_only=True)
        result.page_count = len(wb.sheetnames)
        text_parts = []
        all_tables = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            sheet_table = []
            sheet_text = [f"[Sheet: {sheet_name}]"]

            for row in ws.iter_rows(values_only=True):
                row_vals = [str(c).strip() if c is not None else "" for c in row]
                non_empty = [v for v in row_vals if v]
                if non_empty:
                    sheet_table.append(row_vals)
                    sheet_text.append("  |  ".join(non_empty))

            if sheet_table:
                all_tables.append(sheet_table)
            text_parts.extend(sheet_text)

        result.full_text = "\n".join(text_parts)
        result.raw_tables = all_tables

    # ── CSV ───────────────────────────────────────────────────────────────── 

    async def _ingest_csv(self, result: IngestionResult):
        """Parse CSV into table and text."""
        import csv
        import chardet

        with open(self.file_path, "rb") as f:
            raw = f.read()

        encoding = chardet.detect(raw).get("encoding") or "utf-8"
        text_content = raw.decode(encoding, errors="replace")

        reader = csv.reader(io.StringIO(text_content))
        table = [row for row in reader if any(c.strip() for c in row)]

        result.raw_tables = [table]
        result.full_text = "\n".join([", ".join(row) for row in table])
        result.page_count = 1

    # ── Word ──────────────────────────────────────────────────────────────── 

    async def _ingest_word(self, result: IngestionResult):
        """Extract text and tables from DOCX."""
        from docx import Document

        doc = Document(self.file_path)
        text_parts = [para.text for para in doc.paragraphs if para.text.strip()]

        all_tables = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                rows.append([cell.text.strip() for cell in row.cells])
            if rows:
                all_tables.append(rows)

        result.full_text = "\n".join(text_parts)
        result.raw_tables = all_tables
        result.page_count = 1

    # ── Text ──────────────────────────────────────────────────────────────── 

    async def _ingest_text(self, result: IngestionResult):
        """Read plain text files with encoding detection."""
        import chardet

        with open(self.file_path, "rb") as f:
            raw = f.read()

        encoding = chardet.detect(raw).get("encoding") or "utf-8"
        result.full_text = raw.decode(encoding, errors="replace")
        result.page_count = 1

    # ── ZIP ───────────────────────────────────────────────────────────────── 

    async def _ingest_zip(self, result: IngestionResult):
        """Extract ZIP and ingest each file, concatenating all text."""
        logger.info(f"Extracting ZIP archive: {self.path.name}")
        all_texts = []
        all_tables = []
        total_pages = 0

        with tempfile.TemporaryDirectory() as tmp_dir:
            try:
                with zipfile.ZipFile(self.file_path, "r") as zf:
                    zf.extractall(tmp_dir)
            except zipfile.BadZipFile:
                logger.error("Bad ZIP file — cannot extract")
                result.full_text = ""
                return

            for root, dirs, files in os.walk(tmp_dir):
                for fname in sorted(files):
                    fpath = os.path.join(root, fname)
                    ext = Path(fname).suffix.lower()

                    if ext not in ALL_SUPPORTED:
                        continue

                    try:
                        sub_engine = DocumentIngestionEngine(
                            fpath,
                            ocr_enabled=self.ocr_enabled,
                            preprocess=self.preprocess,
                        )
                        sub_result = await sub_engine.ingest()
                        all_texts.append(f"\n\n=== File: {fname} ===\n{sub_result.full_text}")
                        all_tables.extend(sub_result.raw_tables)
                        total_pages += sub_result.page_count
                        if sub_result.is_ocr_used:
                            result.is_ocr_used = True
                    except Exception as e:
                        logger.warning(f"Skipping file {fname} in ZIP: {e}")

        result.full_text = "\n".join(all_texts)
        result.raw_tables = all_tables
        result.page_count = total_pages


# ── OCR helper ──────────────────────────────────────────────────────────────

def _run_ocr(image_bytes: bytes, lang: str = "eng+hin") -> str:
    """Run Tesseract OCR on preprocessed image bytes."""
    try:
        import pytesseract
        from PIL import Image

        image = Image.open(io.BytesIO(image_bytes))
        config = "--oem 3 --psm 6"
        text = pytesseract.image_to_string(image, lang=lang, config=config)
        return text
    except Exception as e:
        logger.error(f"OCR failed: {e}")
        return ""


def _convert_heic_to_png(heic_bytes: bytes) -> bytes:
    """Convert HEIC image bytes to PNG bytes."""
    try:
        from PIL import Image
        import pillow_heif
        pillow_heif.register_heif_opener()
        img = Image.open(io.BytesIO(heic_bytes))
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        return buf.getvalue()
    except Exception as e:
        logger.warning(f"HEIC conversion failed: {e} — returning raw bytes")
        return heic_bytes
